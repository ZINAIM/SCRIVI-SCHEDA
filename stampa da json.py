from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import json
import os

def stampa_scheda_da_json(file_json, settimane, output_doc="scheda_finale.docx"):
    # === Carica dati JSON ===
    with open(file_json, "r", encoding="utf-8") as f:
        scheda_dati = json.load(f)

    # === Crea nuovo documento Word ===
    doc = Document()
    doc.add_heading("Scheda Allenamento", 0)

    # === Per ogni settimana ===
    for sett in range(settimane):
        doc.add_heading(f"\nSettimana {sett + 1}", level=1)

        for giorno, esercizi in scheda_dati.items():
            doc.add_heading(giorno, level=2)

            for esercizio in esercizi:
                nome = esercizio["esercizio"]

                # Estrai serie/ripetizioni solo se disponibili per questa settimana
                serie_list = esercizio.get("serie", [])
                rip_list = esercizio.get("ripetizioni", [])
                recupero = esercizio.get("recupero", "")

                serie = serie_list[sett] if sett < len(serie_list) else ""
                ripetizioni = rip_list[sett] if sett < len(rip_list) else ""

                if serie and ripetizioni:
                    testo = f"{nome:<40} {serie} x {ripetizioni}   rec {recupero}"
                elif serie:
                    testo = f"{nome:<40} {serie} x   rec {recupero}"
                elif ripetizioni:
                    testo = f"{nome:<40} x {ripetizioni}   rec {recupero}"
                else:
                    testo = f"{nome:<40}"  # Solo il nome

                p = doc.add_paragraph()
                run = p.add_run(testo)
                run.font.name = 'Courier New'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')
                run.font.size = Pt(11)

    # === Salva documento Word ===
    doc.save(output_doc)
    print(f"\n✅ Scheda salvata come '{output_doc}'")

if __name__== "__main__":
    stampa_scheda_da_json("scheda_temp.json", settimane=4, output_doc="scheda_finale.docx")