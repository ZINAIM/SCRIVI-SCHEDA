import json
from collections import OrderedDict
from prompt_toolkit import prompt
from prompt_toolkit.completion import WordCompleter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from aggiorna_link import aggiorna_link
from docx.oxml.ns import qn
import os


# === CONFIG FILE PATH ===
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LINKS_JSON = os.path.join(BASE_DIR, "video_links.json")
MANCANTI_TXT = os.path.join(BASE_DIR, "mancanti.txt")
OUTPUT_DOC = os.path.join(BASE_DIR, "prova scheda.docx")


# === Carica link esercizi da JSON ===   
try:
    with open(LINKS_JSON, "r", encoding="utf-8") as f:
        video_links = json.load(f)
except FileNotFoundError:
    video_links = {}

# === Autocompletamento da elenco esercizi noti ===
esercizi_noti = list(video_links.keys())
esercizio_completer = WordCompleter(esercizi_noti, ignore_case=True)


# === Input utente ===
data = input("Data di oggi: ")
nome_cliente = input("Nome e cognome: ")
giorni_allenamento = input("Quanti giorni ti allenerai? ")

lista_scheda = []
for giorno in range(int(giorni_allenamento)):
    print(f"\nInserisci gli esercizi per il Giorno {chr(65 + giorno)}:")
    esercizi_giorno = []
    while True:
        esercizio = prompt("Nome esercizio (s per stop): ", completer=esercizio_completer)
        if esercizio.lower() == "s":
            break
        serie = input("Numero di serie: ")
        ripetizioni = input("Numero di ripetizioni: ")
        recupero = input("Recupero: ")
        esercizi_giorno.append({
            "esercizio": esercizio,
            "serie": serie,
            "ripetizioni": ripetizioni,
            "recupero": recupero
        })
    lista_scheda.append(esercizi_giorno)

# === Crea documento Word ===
scheda = Document()
scheda.add_heading(f"Scheda {data}, {nome_cliente}", 0)

# Trova la lunghezza max per allineamento
max_len = max((len(e["esercizio"]) for giorno in lista_scheda for e in giorno), default=0)

# Aggiungi giorni con esercizi
for giorno_idx, giorno in enumerate(lista_scheda):
    scheda.add_heading(f"\n Giorno {chr(65 + giorno_idx)}", 2)
    for esercizio in giorno:
        p = scheda.add_paragraph()
        run = p.add_run(f"{esercizio['esercizio']:<{max_len+2}} {esercizio['serie']} x {esercizio['ripetizioni']}   rec {esercizio['recupero']}'")
        run.font.name = 'Courier New'
        p.paragraph_format.space_after = Pt(0)

# === Elenco esercizi finali con link video ===
esercizi_unici = OrderedDict()
for giorno in lista_scheda:
    for esercizio in giorno:
        nome = esercizio["esercizio"]
        esercizi_unici[nome] = None  # mantiene l'ordine

scheda.add_page_break()
scheda.add_heading("Video esercizi", level=1)

mancanti = []

for nome in esercizi_unici:
    link = video_links.get(nome)
    if link:
        testo = f"{nome}: {link}"
    else:
        testo = f"{nome}: 🔗 [manca link]"
        mancanti.append(nome)
    p = scheda.add_paragraph()
    run = p.add_run(testo)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')  

# === Salva documento finale ===
scheda.save(OUTPUT_DOC)

# === Scrivi esercizi mancanti da completare ===
# Dopo aver creato la lista mancanti (che hai già nel tuo script)

if mancanti:
    with open(MANCANTI_TXT, "w", encoding="utf-8") as f:
        for nome in mancanti:
            f.write(nome + "\n")    
    risposta = input(f"\nCi sono {len(mancanti)} esercizi senza link. Vuoi aggiornarli adesso? (s/n): ").strip().lower()
    if risposta == "s":
        aggiorna_link()  # usa i file di default o passa i nomi se vuoi
    else:
        print("Aggiornamento rimandato. Puoi farlo più tardi con lo script dedicato.")
else:
    print("\nTutti gli esercizi hanno un link.")
