import json
from collections import OrderedDict
from prompt_toolkit import prompt
from prompt_toolkit.completion import WordCompleter
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from aggiorna_link import aggiorna_link_da_lista
from docx.oxml.ns import qn
import os
import webbrowser

# === CONFIG FILE PATH ===
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
LINKS_JSON = os.path.join(BASE_DIR, "video_links.json")
MANCANTI_TXT = os.path.join(BASE_DIR, "mancanti.txt")
OUTPUT_DOC = os.path.join(BASE_DIR, "prova scheda.docx")
TEMP_JSON = os.path.join(BASE_DIR, "scheda_temp.json")

# === Carica link esercizi da JSON ===
try:
    with open(LINKS_JSON, "r", encoding="utf-8") as f:
        video_links = json.load(f)
except FileNotFoundError:
    video_links = {}

# === Autocompletamento da elenco esercizi noti ===
esercizi_noti = list(video_links.keys())
esercizio_completer = WordCompleter(esercizi_noti, ignore_case=True)

# === Input utente ===
data = input("Data di oggi: ")
nome_cliente = input("Nome e cognome: ")
giorni_allenamento = int(input("Quanti giorni ti allenerai? "))
settimane_allenamento = int(input("Quante settimane ti allenerai? "))

# === Costruisci struttura scheda ===
scheda_dict = {}
for giorno in range(giorni_allenamento):
    giorno_nome = f"Giorno {chr(65 + giorno)}"
    print(f"\nInserisci gli esercizi per il {giorno_nome}:")
    esercizi_giorno = []

    while True:
        esercizio = prompt("Nome esercizio (s per stop): ", completer=esercizio_completer)
        if esercizio.lower() == "s":
            break
        serie = input("Numero di serie: ")
        ripetizioni = input("Numero di ripetizioni: ")
        recupero = input("Recupero: ")

        esercizi_giorno.append({
            "esercizio": esercizio,
            "serie": [serie] * settimane_allenamento,
            "ripetizioni": [ripetizioni] * settimane_allenamento,
            "recupero": recupero
        })

    scheda_dict[giorno_nome] = esercizi_giorno

# === Salva JSON temporaneo per modifica manuale ===
with open(TEMP_JSON, "w", encoding="utf-8") as f:
    json.dump(scheda_dict, f, ensure_ascii=False, indent=2)

print(f"\nFile temporaneo salvato in '{TEMP_JSON}'. Puoi modificarlo ora.")
webbrowser.open(TEMP_JSON)
input("\nPremi INVIO quando hai finito di modificare il file JSON...")

# === Ricarica struttura modificata ===
with open(TEMP_JSON, "r", encoding="utf-8") as f:
    scheda_dict = json.load(f)

# === Crea documento Word ===
scheda = Document()
scheda.add_heading(f"Scheda {data}, {nome_cliente}", 0)

# === Stampa settimana per settimana ===
esercizi_unici = OrderedDict()
for settimana in range(settimane_allenamento):
    scheda.add_heading(f"Settimana {settimana + 1}", level=1)
    for giorno_nome, esercizi in scheda_dict.items():
        scheda.add_heading(giorno_nome, level=2)
        for esercizio in esercizi:
            nome = esercizio["esercizio"]
            
            # Estrai serie, ripetizioni e recupero per la settimana corrente
            serie_set = esercizio.get("serie", [])
            rip_set = esercizio.get("ripetizioni", [])
            rec = esercizio.get("recupero", "").strip()

            serie = serie_set[settimana] if settimana < len(serie_set) else ""
            rip = rip_set[settimana] if settimana < len(rip_set) else ""

            # Costruzione testo condizionale
            dettaglio = ""
            if serie and rip:
                dettaglio += f"{serie} x {rip}"
            elif serie:
                dettaglio += f"{serie} x"
            elif rip:
                dettaglio += f"x {rip}"
            if rec:
                dettaglio += f"   rec {rec}"

            # Scrivi nel documento, giustificato a destra
            p = scheda.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            run = p.add_run(f"{nome} {dettaglio}".strip())
            run.font.name = 'Courier New'
            esercizi_unici[nome] = None


# === Sezione video ===
scheda.add_page_break()
scheda.add_heading("Video esercizi", level=1)

mancanti = []
for nome in esercizi_unici:
    link = video_links.get(nome)
    testo = f"{nome}: {link}" if link else f"{nome}: 🔗 [manca link]"
    if not link:
        mancanti.append(nome)
    p = scheda.add_paragraph()
    run = p.add_run(testo)
    run.font.name = 'Courier New'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Courier New')

# === Salva documento ===
scheda.save(OUTPUT_DOC)
print(f"\nScheda salvata in '{OUTPUT_DOC}'.")

# === Aggiorna link se mancano ===
if mancanti:
    with open(MANCANTI_TXT, "w", encoding="utf-8") as f:
        for nome in mancanti:
            f.write(nome + "\n")
    risposta = input(f"\nCi sono {len(mancanti)} esercizi senza link. Vuoi aggiornarli adesso? (s/n): ").strip().lower()
    if risposta == "s":
                # Forza apertura e chiusura del file per rilasciare eventuali lock
        try:
            with open(LINKS_JSON, "r", encoding="utf-8") as f:
                pass
        except Exception as e:
            print(f"Errore aprendo file prima aggiornamento link: {e}")

        aggiorna_link_da_lista(mancanti)

# Ora carica i link aggiornati
with open(LINKS_JSON, "r", encoding="utf-8") as f:
    video_links = json.load(f)

input("\nPremi INVIO per chiudere...")
